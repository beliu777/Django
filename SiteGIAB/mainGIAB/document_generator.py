from datetime import datetime, timedelta
import os
from urllib.parse import quote

from django.conf import settings
from django.http import HttpResponse
from docx import Document as DocxDocument

from .models import WordTemplate

# Папка, где лежат старые Word-шаблоны: SiteGIAB/templates/documents/
TEMPLATES_DIR = os.path.join(settings.BASE_DIR, 'templates', 'documents')


def _value(data, field_name):
    return str(getattr(data, field_name, '') or '')


def _int_value(data, field_name):
    value = getattr(data, field_name, None)
    return value if value is not None else ''


def _parse_date(value):
    """Понимает даты из формы: 25.04.2026, 2026-04-25, 25/04/2026."""
    if not value:
        return None
    value = str(value).strip()
    for fmt in ('%d.%m.%Y', '%Y-%m-%d', '%d/%m/%Y'):
        try:
            return datetime.strptime(value, fmt).date()
        except ValueError:
            pass
    return None


def _practice_days(data):
    """
    Список дней практики для циклов в Word-шаблоне.

    Пример в таблице Word:
    {% for day in practice_days %}
    {{ day.number }} | {{ day.date }} | {{ day.weekday }} | {{ day.task }}
    {% endfor %}
    """
    start = _parse_date(getattr(data, 'date_begin', None))
    finish = _parse_date(getattr(data, 'date_finish', None))
    if not start or not finish or finish < start:
        return []

    weekdays = ['понедельник', 'вторник', 'среда', 'четверг', 'пятница', 'суббота', 'воскресенье']
    days = []
    current = start
    number = 1
    while current <= finish:
        days.append({
            'number': number,
            'date': current.strftime('%d.%m.%Y'),
            'weekday': weekdays[current.weekday()],
            'task': '',
        })
        current += timedelta(days=1)
        number += 1
    return days


def _template_context(data):
    """Контекст для обычных {{ field }} и умных Jinja-блоков в Word."""
    full_name = ' '.join(filter(None, [_value(data, 'familia'), _value(data, 'name'), _value(data, 'otchestvo')]))
    initials = ''
    if _value(data, 'name'):
        initials += _value(data, 'name')[0] + '.'
    if _value(data, 'otchestvo'):
        initials += _value(data, 'otchestvo')[0] + '.'
    short_name = ' '.join(filter(None, [_value(data, 'familia'), initials]))

    context = {
        'familia': _value(data, 'familia'),
        'name': _value(data, 'name'),
        'otchestvo': _value(data, 'otchestvo'),
        'fio': full_name,
        'full_name': full_name,
        'short_name': short_name,
        'fio_genitive': _value(data, 'fio_genitive'),
        'tip': _value(data, 'tip'),
        'tip_title': _value(data, 'tip_title'),
        'module': _value(data, 'module'),
        'module_code': _value(data, 'module_code'),
        'specialization': _value(data, 'specialization'),
        'kurs': _int_value(data, 'kurs'),
        'group': _value(data, 'group'),
        'date_begin': _value(data, 'date_begin'),
        'date_finish': _value(data, 'date_finish'),
        'head1': _value(data, 'head1'),
        'head2': _value(data, 'head2'),
        'ruc_pract': _value(data, 'ruc_pract'),
        'year': _int_value(data, 'year'),
        'practice_days': _practice_days(data),
    }
    context['practice_period'] = '{} - {}'.format(context['date_begin'], context['date_finish']).strip(' -')
    context['has_head2'] = bool(context['head2'])
    context['has_ruc_pract'] = bool(context['ruc_pract'])
    context['has_practice_days'] = bool(context['practice_days'])
    return context


def _find_template(possible_names, contains_words, template_type=None):
    """
    Ищет шаблон.
    1) Сначала берет файл, загруженный через Django-админку.
    2) Если в админке файла нет, ищет старые шаблоны в папке templates/documents.
    """
    if template_type:
        uploaded = WordTemplate.objects.filter(template_type=template_type).first()
        if uploaded and uploaded.file and os.path.exists(uploaded.file.path):
            return uploaded.file.path

    for filename in possible_names:
        path = os.path.join(TEMPLATES_DIR, filename)
        if os.path.exists(path):
            return path

    if os.path.isdir(TEMPLATES_DIR):
        for filename in os.listdir(TEMPLATES_DIR):
            low = filename.lower()
            if low.endswith('.docx') and not low.startswith('~$'):
                if any(word.lower() in low for word in contains_words):
                    return os.path.join(TEMPLATES_DIR, filename)

    existing = []
    if os.path.isdir(TEMPLATES_DIR):
        existing = [name for name in os.listdir(TEMPLATES_DIR) if name.lower().endswith('.docx')]

    raise FileNotFoundError(
        'Не найден Word-шаблон. Загрузите его в админке: /admin/ -> Word-шаблоны. '
        'Или положите файл в папку: {}. Искомые имена: {}. Сейчас в папке: {}'.format(
            TEMPLATES_DIR,
            ', '.join(possible_names),
            ', '.join(existing) if existing else 'нет .docx файлов'
        )
    )


def _replace_in_paragraph(paragraph, replacements):
    """Старый режим: заменяет простые {{ поле }}, даже если Word разбил текст на runs."""
    if not paragraph.runs:
        return

    full_text = ''.join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, value in replacements.items():
        if isinstance(value, (list, dict)):
            continue
        value = str(value)
        new_text = new_text.replace('{{ ' + key + ' }}', value)
        new_text = new_text.replace('{{' + key + '}}', value)

    if new_text != full_text:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def _replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacements)


def _fill_with_python_docx(template_path, context):
    """Запасной режим, если docxtpl не установлен: простые {{ field }} без условий/циклов."""
    doc = DocxDocument(template_path)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, context)

    for table in doc.tables:
        _replace_in_table(table, context)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, context)
        for table in section.header.tables:
            _replace_in_table(table, context)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, context)
        for table in section.footer.tables:
            _replace_in_table(table, context)

    return doc


def _fill_docx_template(template_path, output_filename, context):
    """
    Новый умный режим через docxtpl:
    - обычные поля: {{ fio }}, {{ group }}
    - условия: {% if head2 %} ... {% endif %}
    - циклы: {% for day in practice_days %} ... {% endfor %}

    Если docxtpl вдруг не установлен, сайт не падает и продолжает работать со старыми простыми полями.
    """
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = "attachment; filename*=UTF-8''{}".format(quote(output_filename))

    try:
        from docxtpl import DocxTemplate
        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(response)
    except ImportError:
        doc = _fill_with_python_docx(template_path, context)
        doc.save(response)

    return response


def generate_title_page(data):
    template_path = _find_template(
        ['title.docx', 'Титульный.docx', 'Титульный лист.docx', 'Титульник.docx'],
        ['титул'],
        WordTemplate.TITLE,
    )
    return _fill_docx_template(template_path, 'titulny_list.docx', _template_context(data))


def generate_assignment(data):
    template_path = _find_template(
        ['assignment.docx', 'Задание.docx', 'Индивидуальное задание.docx'],
        ['задани'],
        WordTemplate.ASSIGNMENT,
    )
    return _fill_docx_template(template_path, 'zadanie.docx', _template_context(data))


def generate_diary(data):
    template_path = _find_template(
        ['diary.docx', 'Дневник.docx', 'Дневник_f7lmcKg.docx'],
        ['дневник'],
        WordTemplate.DIARY,
    )
    return _fill_docx_template(template_path, 'dnevnik.docx', _template_context(data))
